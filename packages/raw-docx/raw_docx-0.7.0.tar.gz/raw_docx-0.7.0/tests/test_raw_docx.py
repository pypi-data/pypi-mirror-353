import pytest
import os
from docx import Document as DocxDocument
from docx.shared import Inches
from src.raw_docx.raw_docx import RawDocx
from src.raw_docx.raw_document import RawDocument


@pytest.fixture
def raw_docx():
    return RawDocx("tests/test_files/example_1.docx")


@pytest.fixture
def temp_docx(tmp_path):
    """Create a test docx file with various content"""
    doc_path = tmp_path / "test.docx"
    doc = DocxDocument()

    # Add regular paragraph
    doc.add_paragraph("Regular paragraph")

    # Add list items using standard list style
    doc.add_paragraph("First bullet point", style="List Bullet")
    doc.add_paragraph("Second bullet point", style="List Bullet")

    # Add table with merged cells
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))  # Merge first row
    table.cell(0, 0).text = "Merged cell"
    table.cell(1, 0).text = "Cell 1"
    table.cell(1, 1).text = "Cell 2"

    # Add image if test file exists
    image_path = os.path.join(os.path.dirname(__file__), "test_files", "test_image.png")
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(1.0))

    # Save the document
    doc.save(doc_path)
    return str(doc_path)


def test_to_dict_with_document(raw_docx):
    """Test converting RawDocx to dictionary with loaded document"""
    result = raw_docx.to_dict()
    assert result["type"] == "raw_docx"
    assert result["document"] is not None
    assert result["document"]["type"] == "document"
    assert isinstance(result["document"]["sections"], list)


def test_initialization_and_processing(temp_docx):
    """Test document initialization and processing"""
    docx = RawDocx(temp_docx)
    assert os.path.exists(docx.image_path)
    assert isinstance(docx.target_document, RawDocument)
    assert len(docx.target_document.sections) > 0


def test_table_processing(temp_docx):
    """Test table processing with merged cells"""
    docx = RawDocx(temp_docx)
    tables = docx.target_document.sections[0].tables()
    assert len(tables) > 0

    # Check merged cells
    first_table = tables[0]
    first_row = first_table.rows[0]
    assert len(first_row.cells) == 2
    assert first_row.cells[0].h_span == 2  # Horizontally merged
    assert first_row.cells[0].first is True
    assert first_row.cells[1].first is False


def test_image_processing(temp_docx):
    """Test image extraction and processing"""
    docx = RawDocx(temp_docx)

    # Check if image directory was created
    assert os.path.exists(docx.image_path)

    # Check if image is referenced in document
    image_path = os.path.join(os.path.dirname(__file__), "test_files", "test_image.png")
    if os.path.exists(image_path):
        found_image = False
        for section in docx.target_document.sections:
            for item in section.items:
                if hasattr(item, "filepath") and item.filepath.endswith(
                    (".png", ".jpg", ".jpeg")
                ):
                    found_image = True
                    break
        assert found_image, "Image not found in document"


def test_error_handling(tmp_path):
    """Test error handling for invalid files and directories"""
    # Test with non-existent file
    with pytest.raises(Exception):
        RawDocx(str(tmp_path / "nonexistent.docx"))

    # Test with invalid file format
    invalid_file = tmp_path / "invalid.txt"
    invalid_file.write_text("Not a docx file")
    with pytest.raises(Exception):
        RawDocx(str(invalid_file))
