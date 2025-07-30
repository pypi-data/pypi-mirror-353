import pytest
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from src.raw_docx.raw_run import RawRun


@pytest.fixture
def mock_docx_paragraph():
    """Create a test paragraph with multiple runs"""
    doc = Document()
    paragraph = doc.add_paragraph()

    # Add normal run
    run1 = paragraph.add_run("Hello ")
    run1.font.color.rgb = RGBColor(255, 0, 0)  # Red

    # Add styled run
    run2 = paragraph.add_run("World")
    run2.font.color.rgb = RGBColor(0, 0, 255)  # Blue

    # Add highlighted run
    run3 = paragraph.add_run("!")
    run3.font.highlight_color = WD_COLOR_INDEX.YELLOW

    return paragraph


def test_raw_run_initialization():
    """Test RawRun initialization with basic attributes and strip whitespace"""
    text = " Hello World "
    color = "FF0000"  # Red
    highlight = "yellow"
    style = "Normal"

    run = RawRun(text=text, color=color, highlight=highlight, style=style)

    assert run.text == " Hello World "
    assert run.color == color
    assert run.highlight == highlight
    assert run.style == style


def test_raw_run_initialization_empty_values():
    """Test RawRun initialization with empty values"""
    text = ""
    color = None
    highlight = None
    style = ""

    run = RawRun(text=text, color=color, highlight=highlight, style=style)

    assert run.text == text
    assert run.color == color
    assert run.highlight == highlight
    assert run.style == style


def test_to_dict():
    """Test converting RawRun to dictionary"""
    text = "Hello World"
    color = "FF0000"
    highlight = "yellow"
    style = "Normal"

    run = RawRun(text=text, color=color, highlight=highlight, style=style)
    run_dict = run.to_dict()

    assert run_dict == {
        "text": text,
        "color": color,
        "highlight": highlight,
        "style": style,
    }


def test_to_dict_empty_values():
    """Test converting RawRun with empty values to dictionary"""
    text = ""
    color = None
    highlight = None
    style = ""

    run = RawRun(text=text, color=color, highlight=highlight, style=style)
    run_dict = run.to_dict()

    assert run_dict == {
        "text": text,
        "color": color,
        "highlight": highlight,
        "style": style,
    }
