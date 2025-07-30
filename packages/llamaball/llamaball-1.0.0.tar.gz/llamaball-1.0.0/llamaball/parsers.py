"""
Llamaball - File Parsers Module
File Purpose: Comprehensive file type parsing for document ingestion
Primary Functions: Parse various file formats into text for embedding
Inputs: File paths and content
Outputs: Extracted text content with metadata
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
import json
import csv
from io import StringIO

# Optional imports with fallbacks
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.pdfparser import PDFSyntaxError
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    
try:
    import openpyxl
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    
try:
    import xlrd
    XLS_AVAILABLE = True
except ImportError:
    XLS_AVAILABLE = False

logger = logging.getLogger(__name__)

class FileParser:
    """High-performance file parser with support for multiple formats."""
    
    # Supported file extensions by category
    TEXT_EXTENSIONS = {
        '.txt', '.md', '.rst', '.tex', '.org', '.adoc', '.wiki',
        '.markdown', '.mdown', '.mkd', '.text', '.asc'
    }
    
    CODE_EXTENSIONS = {
        '.py', '.js', '.ts', '.jsx', '.tsx', '.html', '.htm', '.css',
        '.json', '.xml', '.yaml', '.yml', '.toml', '.ini', '.cfg',
        '.sql', '.sh', '.bash', '.zsh', '.fish', '.ps1', '.bat',
        '.php', '.rb', '.go', '.rs', '.cpp', '.c', '.h', '.hpp',
        '.java', '.scala', '.kt', '.swift', '.dart', '.r', '.m',
        '.pl', '.lua', '.vim', '.dockerfile', '.makefile'
    }
    
    DATA_EXTENSIONS = {
        '.csv', '.tsv', '.jsonl', '.ndjson', '.log'
    }
    
    DOCUMENT_EXTENSIONS = {
        '.pdf', '.docx', '.doc'
    }
    
    SPREADSHEET_EXTENSIONS = {
        '.xlsx', '.xls', '.xlsm', '.ods'
    }
    
    NOTEBOOK_EXTENSIONS = {
        '.ipynb'
    }
    
    @classmethod
    def get_supported_extensions(cls) -> set:
        """Get all supported file extensions."""
        return (cls.TEXT_EXTENSIONS | cls.CODE_EXTENSIONS | 
                cls.DATA_EXTENSIONS | cls.DOCUMENT_EXTENSIONS |
                cls.SPREADSHEET_EXTENSIONS | cls.NOTEBOOK_EXTENSIONS)
    
    @classmethod
    def is_supported(cls, file_path: Union[str, Path]) -> bool:
        """Check if file type is supported."""
        ext = Path(file_path).suffix.lower()
        return ext in cls.get_supported_extensions()
    
    @classmethod
    def get_file_type(cls, file_path: Union[str, Path]) -> str:
        """Determine file type category."""
        ext = Path(file_path).suffix.lower()
        
        if ext in cls.TEXT_EXTENSIONS:
            return "text"
        elif ext in cls.CODE_EXTENSIONS:
            return "code"
        elif ext in cls.DATA_EXTENSIONS:
            return "data"
        elif ext in cls.DOCUMENT_EXTENSIONS:
            return "document"
        elif ext in cls.SPREADSHEET_EXTENSIONS:
            return "spreadsheet"
        elif ext in cls.NOTEBOOK_EXTENSIONS:
            return "notebook"
        else:
            return "unknown"
    
    def parse_file(self, file_path: Union[str, Path]) -> Dict[str, Union[str, Dict]]:
        """
        Parse file and extract text content with metadata.
        
        Args:
            file_path: Path to file to parse
            
        Returns:
            Dictionary with 'content', 'metadata', and 'error' keys
            
        Performance:
            - Optimized for large files with streaming when possible
            - Memory-efficient processing for PDFs and spreadsheets
            - Fast text extraction with minimal overhead
        """
        file_path = Path(file_path)
        
        result = {
            'content': '',
            'metadata': {
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size': 0,
                'file_type': self.get_file_type(file_path),
                'extension': file_path.suffix.lower(),
                'encoding': 'utf-8'
            },
            'error': None
        }
        
        try:
            # Get file size
            if file_path.exists():
                result['metadata']['file_size'] = file_path.stat().st_size
            else:
                result['error'] = f"File not found: {file_path}"
                return result
            
            # Parse based on file type
            ext = file_path.suffix.lower()
            
            if ext == '.pdf':
                result.update(self._parse_pdf(file_path))
            elif ext in {'.docx', '.doc'}:
                result.update(self._parse_docx(file_path))
            elif ext in {'.xlsx', '.xls', '.xlsm'}:
                result.update(self._parse_excel(file_path))
            elif ext == '.csv':
                result.update(self._parse_csv(file_path))
            elif ext == '.tsv':
                result.update(self._parse_tsv(file_path))
            elif ext in {'.json', '.jsonl', '.ndjson'}:
                result.update(self._parse_json(file_path))
            elif ext == '.ipynb':
                result.update(self._parse_notebook(file_path))
            elif ext in self.TEXT_EXTENSIONS | self.CODE_EXTENSIONS:
                result.update(self._parse_text(file_path))
            else:
                # Try as text file
                result.update(self._parse_text(file_path))
                
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            result['error'] = str(e)
        
        return result
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse PDF files using pdfminer with robust error handling."""
        if not PDF_AVAILABLE:
            return {
                'content': '',
                'error': 'PDF parsing not available. Install with: pip install pdfminer.six'
            }
        
        try:
            # Extract text with optimized settings and error recovery
            content = pdf_extract_text(
                str(file_path),
                maxpages=0,  # Process all pages
                caching=True,
                codec='utf-8'
            )
            
            # Clean up extracted text
            content = self._clean_text(content)
            
            # Handle empty content
            if not content or len(content.strip()) < 10:
                return {
                    'content': '',
                    'error': 'PDF appears to be empty or contains only images/scanned content'
                }
            
            return {
                'content': content,
                'metadata': {'pages': content.count('\f') + 1 if content else 0}
            }
            
        except PDFSyntaxError as e:
            error_msg = str(e)
            if "No /Root object" in error_msg:
                return {
                    'content': '', 
                    'error': 'PDF file appears to be corrupted or not a valid PDF'
                }
            return {'content': '', 'error': f'PDF syntax error: {error_msg}'}
        except Exception as e:
            error_msg = str(e)
            if "No /Root object" in error_msg:
                return {
                    'content': '', 
                    'error': 'PDF file appears to be corrupted or not a valid PDF'
                }
            elif "password" in error_msg.lower():
                return {
                    'content': '', 
                    'error': 'PDF is password protected'
                }
            elif "encrypted" in error_msg.lower():
                return {
                    'content': '', 
                    'error': 'PDF is encrypted and cannot be parsed'
                }
            return {'content': '', 'error': f'PDF parsing error: {error_msg}'}
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse DOCX files using python-docx."""
        if not DOCX_AVAILABLE:
            return {
                'content': '',
                'error': 'DOCX parsing not available. Install with: pip install python-docx'
            }
        
        try:
            doc = Document(str(file_path))
            
            # Extract paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        tables_text.append(row_text)
            
            # Combine content
            content_parts = paragraphs
            if tables_text:
                content_parts.extend(['', '--- Tables ---'] + tables_text)
            
            content = '\n'.join(content_parts)
            
            return {
                'content': content,
                'metadata': {
                    'paragraphs': len(paragraphs),
                    'tables': len(doc.tables)
                }
            }
            
        except Exception as e:
            return {'content': '', 'error': f'DOCX parsing error: {e}'}
    
    def _parse_excel(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse Excel files (xlsx, xls)."""
        ext = file_path.suffix.lower()
        
        if ext in {'.xlsx', '.xlsm'} and EXCEL_AVAILABLE:
            return self._parse_xlsx(file_path)
        elif ext == '.xls' and XLS_AVAILABLE:
            return self._parse_xls(file_path)
        else:
            return {
                'content': '',
                'error': f'Excel parsing not available for {ext}. Install with: pip install openpyxl xlrd'
            }
    
    def _parse_xlsx(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse XLSX files using openpyxl."""
        try:
            workbook = load_workbook(str(file_path), read_only=True, data_only=True)
            
            sheets_content = []
            total_rows = 0
            
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                
                sheet_content = [f"=== Sheet: {sheet_name} ==="]
                sheet_rows = 0
                
                for row in worksheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                        if row_text.strip():
                            sheet_content.append(row_text)
                            sheet_rows += 1
                
                if sheet_rows > 0:
                    sheets_content.extend(sheet_content + [''])
                    total_rows += sheet_rows
            
            workbook.close()
            
            return {
                'content': '\n'.join(sheets_content),
                'metadata': {
                    'sheets': len(workbook.sheetnames),
                    'total_rows': total_rows
                }
            }
            
        except Exception as e:
            return {'content': '', 'error': f'XLSX parsing error: {e}'}
    
    def _parse_xls(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse XLS files using xlrd."""
        try:
            workbook = xlrd.open_workbook(str(file_path))
            
            sheets_content = []
            total_rows = 0
            
            for sheet_idx in range(workbook.nsheets):
                worksheet = workbook.sheet_by_index(sheet_idx)
                sheet_name = worksheet.name
                
                sheet_content = [f"=== Sheet: {sheet_name} ==="]
                
                for row_idx in range(worksheet.nrows):
                    row = worksheet.row_values(row_idx)
                    row_text = ' | '.join(str(cell) for cell in row if cell)
                    if row_text.strip():
                        sheet_content.append(row_text)
                
                if worksheet.nrows > 0:
                    sheets_content.extend(sheet_content + [''])
                    total_rows += worksheet.nrows
            
            return {
                'content': '\n'.join(sheets_content),
                'metadata': {
                    'sheets': workbook.nsheets,
                    'total_rows': total_rows
                }
            }
            
        except Exception as e:
            return {'content': '', 'error': f'XLS parsing error: {e}'}
    
    def _parse_csv(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse CSV files with automatic delimiter detection."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                # Read sample to detect dialect
                sample = f.read(8192)
                f.seek(0)
                
                # Detect CSV dialect
                try:
                    dialect = csv.Sniffer().sniff(sample)
                    reader = csv.reader(f, dialect)
                except csv.Error:
                    # Fallback to standard CSV
                    f.seek(0)
                    reader = csv.reader(f)
                
                rows = []
                for row_num, row in enumerate(reader):
                    if row_num == 0:
                        # Header row
                        rows.append(' | '.join(f"**{cell}**" for cell in row))
                    else:
                        rows.append(' | '.join(row))
                    
                    # Limit for very large files
                    if row_num > 10000:
                        rows.append(f"... (truncated after {row_num} rows)")
                        break
                
                return {
                    'content': '\n'.join(rows),
                    'metadata': {'rows': len(rows) - 1}  # Exclude header
                }
                
        except Exception as e:
            return {'content': '', 'error': f'CSV parsing error: {e}'}
    
    def _parse_tsv(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse TSV files."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f, delimiter='\t')
                
                rows = []
                for row_num, row in enumerate(reader):
                    if row_num == 0:
                        # Header row
                        rows.append(' | '.join(f"**{cell}**" for cell in row))
                    else:
                        rows.append(' | '.join(row))
                    
                    # Limit for very large files
                    if row_num > 10000:
                        rows.append(f"... (truncated after {row_num} rows)")
                        break
                
                return {
                    'content': '\n'.join(rows),
                    'metadata': {'rows': len(rows) - 1}
                }
                
        except Exception as e:
            return {'content': '', 'error': f'TSV parsing error: {e}'}
    
    def _parse_json(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse JSON and JSONL files."""
        try:
            ext = file_path.suffix.lower()
            
            if ext in {'.jsonl', '.ndjson'}:
                # Line-delimited JSON
                content_lines = []
                with open(file_path, 'r', encoding='utf-8') as f:
                    for line_num, line in enumerate(f):
                        line = line.strip()
                        if line:
                            try:
                                obj = json.loads(line)
                                content_lines.append(json.dumps(obj, indent=2))
                            except json.JSONDecodeError:
                                content_lines.append(f"Invalid JSON on line {line_num + 1}: {line}")
                        
                        # Limit for very large files
                        if line_num > 1000:
                            content_lines.append(f"... (truncated after {line_num} lines)")
                            break
                
                return {
                    'content': '\n---\n'.join(content_lines),
                    'metadata': {'lines': len(content_lines)}
                }
            else:
                # Regular JSON
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Pretty print JSON
                content = json.dumps(data, indent=2, ensure_ascii=False)
                
                return {
                    'content': content,
                    'metadata': {'json_type': type(data).__name__}
                }
                
        except Exception as e:
            return {'content': '', 'error': f'JSON parsing error: {e}'}
    
    def _parse_notebook(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse Jupyter notebook files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                notebook = json.load(f)
            
            content_parts = []
            cell_count = 0
            code_cells = 0
            markdown_cells = 0
            
            for cell in notebook.get('cells', []):
                cell_type = cell.get('cell_type', 'unknown')
                source = cell.get('source', [])
                
                if isinstance(source, list):
                    cell_content = ''.join(source)
                else:
                    cell_content = str(source)
                
                if cell_content.strip():
                    content_parts.append(f"=== {cell_type.title()} Cell ===")
                    content_parts.append(cell_content)
                    content_parts.append('')
                    
                    cell_count += 1
                    if cell_type == 'code':
                        code_cells += 1
                    elif cell_type == 'markdown':
                        markdown_cells += 1
            
            return {
                'content': '\n'.join(content_parts),
                'metadata': {
                    'total_cells': cell_count,
                    'code_cells': code_cells,
                    'markdown_cells': markdown_cells,
                    'kernel': notebook.get('metadata', {}).get('kernelspec', {}).get('name', 'unknown')
                }
            }
            
        except Exception as e:
            return {'content': '', 'error': f'Notebook parsing error: {e}'}
    
    def _parse_text(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse text and code files with encoding detection."""
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()
                
                # Clean and normalize content
                content = self._clean_text(content)
                
                return {
                    'content': content,
                    'metadata': {
                        'encoding': encoding,
                        'lines': content.count('\n') + 1 if content else 0,
                        'characters': len(content)
                    }
                }
                
            except UnicodeDecodeError:
                continue
            except Exception as e:
                return {'content': '', 'error': f'Text parsing error: {e}'}
        
        return {'content': '', 'error': 'Could not decode file with any supported encoding'}
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ''
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive blank lines
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        return text.strip()

# Global parser instance
parser = FileParser()

def parse_file(file_path: Union[str, Path]) -> Dict[str, Union[str, Dict]]:
    """
    Convenience function to parse a file.
    
    Args:
        file_path: Path to file to parse
        
    Returns:
        Dictionary with parsed content and metadata
    """
    return parser.parse_file(file_path)

def get_supported_extensions() -> set:
    """Get all supported file extensions."""
    return FileParser.get_supported_extensions()

def is_supported_file(file_path: Union[str, Path]) -> bool:
    """Check if file type is supported."""
    return FileParser.is_supported(file_path) 