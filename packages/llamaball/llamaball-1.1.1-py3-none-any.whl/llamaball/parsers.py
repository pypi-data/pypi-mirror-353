"""
Llamaball - File Parsers Module
File Purpose: Comprehensive file type parsing for document ingestion
Primary Functions: Parse various file formats into text for embedding
Inputs: File paths and content
Outputs: Extracted text content with metadata
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
import json
import csv
from io import StringIO

# Optional imports with fallbacks
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.pdfparser import PDFSyntaxError
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    
try:
    import openpyxl
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    
try:
    import xlrd
    XLS_AVAILABLE = True
except ImportError:
    XLS_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False

try:
    import email
    from email.parser import BytesParser
    from email.policy import default
    EMAIL_AVAILABLE = True
except ImportError:
    EMAIL_AVAILABLE = False

try:
    import zipfile
    import tarfile
    ARCHIVE_AVAILABLE = True
except ImportError:
    ARCHIVE_AVAILABLE = False

logger = logging.getLogger(__name__)

class FileParser:
    """High-performance file parser with support for multiple formats."""
    
    # Supported file extensions by category
    TEXT_EXTENSIONS = {
        '.txt', '.md', '.rst', '.tex', '.org', '.adoc', '.wiki',
        '.markdown', '.mdown', '.mkd', '.text', '.asc', '.rtf',
        '.textile', '.mediawiki', '.creole', '.bbcode'
    }
    
    CODE_EXTENSIONS = {
        '.py', '.js', '.ts', '.jsx', '.tsx', '.html', '.htm', '.css',
        '.json', '.xml', '.yaml', '.yml', '.toml', '.ini', '.cfg',
        '.sql', '.sh', '.bash', '.zsh', '.fish', '.ps1', '.bat',
        '.php', '.rb', '.go', '.rs', '.cpp', '.c', '.h', '.hpp',
        '.java', '.scala', '.kt', '.swift', '.dart', '.r', '.m',
        '.pl', '.lua', '.vim', '.dockerfile', '.makefile', '.gradle',
        '.cmake', '.scss', '.sass', '.less', '.styl', '.vue', '.svelte',
        '.astro', '.jsx', '.tsx', '.mjs', '.cjs', '.coffee', '.litcoffee'
    }
    
    WEB_EXTENSIONS = {
        '.html', '.htm', '.xhtml', '.xml', '.rss', '.atom', '.svg',
        '.wml', '.xsl', '.xslt', '.jsp', '.asp', '.aspx', '.php'
    }
    
    DATA_EXTENSIONS = {
        '.csv', '.tsv', '.jsonl', '.ndjson', '.log', '.parquet',
        '.arrow', '.feather', '.pickle', '.pkl', '.hdf5', '.h5'
    }
    
    DOCUMENT_EXTENSIONS = {
        '.pdf', '.docx', '.doc', '.odt', '.pages', '.rtf'
    }
    
    SPREADSHEET_EXTENSIONS = {
        '.xlsx', '.xls', '.xlsm', '.ods', '.numbers', '.gnumeric'
    }
    
    NOTEBOOK_EXTENSIONS = {
        '.ipynb', '.rmd', '.qmd', '.rmarkdown'
    }
    
    EMAIL_EXTENSIONS = {
        '.eml', '.msg', '.mbox', '.maildir'
    }
    
    ARCHIVE_EXTENSIONS = {
        '.zip', '.tar', '.tar.gz', '.tgz', '.tar.bz2', '.tbz2',
        '.tar.xz', '.txz', '.rar', '.7z'
    }
    
    CONFIG_EXTENSIONS = {
        '.conf', '.config', '.properties', '.env', '.editorconfig',
        '.gitignore', '.dockerignore', '.eslintrc', '.prettierrc',
        '.babelrc', '.tsconfig', '.package', '.lock'
    }
    
    @classmethod
    def get_supported_extensions(cls) -> set:
        """Get all supported file extensions."""
        return (cls.TEXT_EXTENSIONS | cls.CODE_EXTENSIONS | cls.WEB_EXTENSIONS |
                cls.DATA_EXTENSIONS | cls.DOCUMENT_EXTENSIONS |
                cls.SPREADSHEET_EXTENSIONS | cls.NOTEBOOK_EXTENSIONS |
                cls.EMAIL_EXTENSIONS | cls.ARCHIVE_EXTENSIONS | cls.CONFIG_EXTENSIONS)
    
    @classmethod
    def is_supported(cls, file_path: Union[str, Path]) -> bool:
        """Check if file type is supported."""
        ext = Path(file_path).suffix.lower()
        return ext in cls.get_supported_extensions()
    
    @classmethod
    def get_file_type(cls, file_path: Union[str, Path]) -> str:
        """Determine file type category."""
        ext = Path(file_path).suffix.lower()
        
        if ext in cls.TEXT_EXTENSIONS:
            return "text"
        elif ext in cls.CODE_EXTENSIONS:
            return "code"
        elif ext in cls.WEB_EXTENSIONS:
            return "web"
        elif ext in cls.DATA_EXTENSIONS:
            return "data"
        elif ext in cls.DOCUMENT_EXTENSIONS:
            return "document"
        elif ext in cls.SPREADSHEET_EXTENSIONS:
            return "spreadsheet"
        elif ext in cls.NOTEBOOK_EXTENSIONS:
            return "notebook"
        elif ext in cls.EMAIL_EXTENSIONS:
            return "email"
        elif ext in cls.ARCHIVE_EXTENSIONS:
            return "archive"
        elif ext in cls.CONFIG_EXTENSIONS:
            return "config"
        else:
            return "unknown"
    
    def parse_file(self, file_path: Union[str, Path]) -> Dict[str, Union[str, Dict]]:
        """
        Parse file and extract text content with metadata.
        
        Args:
            file_path: Path to file to parse
            
        Returns:
            Dictionary with 'content', 'metadata', and 'error' keys
            
        Performance:
            - Optimized for large files with streaming when possible
            - Memory-efficient processing for PDFs and spreadsheets
            - Fast text extraction with minimal overhead
        """
        file_path = Path(file_path)
        
        result = {
            'content': '',
            'metadata': {
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size': 0,
                'file_type': self.get_file_type(file_path),
                'extension': file_path.suffix.lower(),
                'encoding': 'utf-8'
            },
            'error': None
        }
        
        try:
            # Get file size
            if file_path.exists():
                result['metadata']['file_size'] = file_path.stat().st_size
            else:
                result['error'] = f"File not found: {file_path}"
                return result
            
            # Parse based on file type
            ext = file_path.suffix.lower()
            
            if ext == '.pdf':
                result.update(self._parse_pdf(file_path))
            elif ext in {'.docx', '.doc'}:
                result.update(self._parse_docx(file_path))
            elif ext in {'.xlsx', '.xls', '.xlsm'}:
                result.update(self._parse_excel(file_path))
            elif ext == '.csv':
                result.update(self._parse_csv(file_path))
            elif ext == '.tsv':
                result.update(self._parse_tsv(file_path))
            elif ext in {'.json', '.jsonl', '.ndjson'}:
                result.update(self._parse_json(file_path))
            elif ext == '.ipynb':
                result.update(self._parse_notebook(file_path))
            elif ext in {'.html', '.htm', '.xhtml'}:
                result.update(self._parse_html(file_path))
            elif ext in {'.xml', '.rss', '.atom', '.svg'}:
                result.update(self._parse_xml(file_path))
            elif ext in {'.eml', '.msg'}:
                result.update(self._parse_email(file_path))
            elif ext == '.rtf':
                result.update(self._parse_rtf(file_path))
            elif ext in self.ARCHIVE_EXTENSIONS:
                result.update(self._parse_archive(file_path))
            elif ext in self.TEXT_EXTENSIONS | self.CODE_EXTENSIONS | self.CONFIG_EXTENSIONS:
                result.update(self._parse_text(file_path))
            else:
                # Try as text file
                result.update(self._parse_text(file_path))
                
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            result['error'] = str(e)
        
        return result
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse PDF files using pdfminer with robust error handling."""
        if not PDF_AVAILABLE:
            return {
                'content': '',
                'error': 'PDF parsing not available. Install with: pip install pdfminer.six'
            }
        
        try:
            # Extract text with optimized settings and error recovery
            content = pdf_extract_text(
                str(file_path),
                maxpages=0,  # Process all pages
                caching=True,
                codec='utf-8'
            )
            
            # Clean up extracted text
            content = self._clean_text(content)
            
            # Handle empty content
            if not content or len(content.strip()) < 10:
                return {
                    'content': '',
                    'error': 'PDF appears to be empty or contains only images/scanned content'
                }
            
            return {
                'content': content,
                'metadata': {'pages': content.count('\f') + 1 if content else 0}
            }
            
        except PDFSyntaxError as e:
            error_msg = str(e)
            if "No /Root object" in error_msg:
                return {
                    'content': '', 
                    'error': 'PDF file appears to be corrupted or not a valid PDF'
                }
            return {'content': '', 'error': f'PDF syntax error: {error_msg}'}
        except Exception as e:
            error_msg = str(e)
            if "No /Root object" in error_msg:
                return {
                    'content': '', 
                    'error': 'PDF file appears to be corrupted or not a valid PDF'
                }
            elif "password" in error_msg.lower():
                return {
                    'content': '', 
                    'error': 'PDF is password protected'
                }
            elif "encrypted" in error_msg.lower():
                return {
                    'content': '', 
                    'error': 'PDF is encrypted and cannot be parsed'
                }
            return {'content': '', 'error': f'PDF parsing error: {error_msg}'}
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse DOCX files using python-docx."""
        if not DOCX_AVAILABLE:
            return {
                'content': '',
                'error': 'DOCX parsing not available. Install with: pip install python-docx'
            }
        
        try:
            doc = Document(str(file_path))
            
            # Extract paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        tables_text.append(row_text)
            
            # Combine content
            content_parts = paragraphs
            if tables_text:
                content_parts.extend(['', '--- Tables ---'] + tables_text)
            
            content = '\n'.join(content_parts)
            
            return {
                'content': content,
                'metadata': {
                    'paragraphs': len(paragraphs),
                    'tables': len(doc.tables)
                }
            }
            
        except Exception as e:
            return {'content': '', 'error': f'DOCX parsing error: {e}'}
    
    def _parse_excel(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse Excel files (xlsx, xls)."""
        ext = file_path.suffix.lower()
        
        if ext in {'.xlsx', '.xlsm'} and EXCEL_AVAILABLE:
            return self._parse_xlsx(file_path)
        elif ext == '.xls' and XLS_AVAILABLE:
            return self._parse_xls(file_path)
        else:
            return {
                'content': '',
                'error': f'Excel parsing not available for {ext}. Install with: pip install openpyxl xlrd'
            }
    
    def _parse_xlsx(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse XLSX files using openpyxl."""
        try:
            workbook = load_workbook(str(file_path), read_only=True, data_only=True)
            
            sheets_content = []
            total_rows = 0
            
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                
                sheet_content = [f"=== Sheet: {sheet_name} ==="]
                sheet_rows = 0
                
                for row in worksheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                        if row_text.strip():
                            sheet_content.append(row_text)
                            sheet_rows += 1
                
                if sheet_rows > 0:
                    sheets_content.extend(sheet_content + [''])
                    total_rows += sheet_rows
            
            workbook.close()
            
            return {
                'content': '\n'.join(sheets_content),
                'metadata': {
                    'sheets': len(workbook.sheetnames),
                    'total_rows': total_rows
                }
            }
            
        except Exception as e:
            return {'content': '', 'error': f'XLSX parsing error: {e}'}
    
    def _parse_xls(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse XLS files using xlrd."""
        try:
            workbook = xlrd.open_workbook(str(file_path))
            
            sheets_content = []
            total_rows = 0
            
            for sheet_idx in range(workbook.nsheets):
                worksheet = workbook.sheet_by_index(sheet_idx)
                sheet_name = worksheet.name
                
                sheet_content = [f"=== Sheet: {sheet_name} ==="]
                
                for row_idx in range(worksheet.nrows):
                    row = worksheet.row_values(row_idx)
                    row_text = ' | '.join(str(cell) for cell in row if cell)
                    if row_text.strip():
                        sheet_content.append(row_text)
                
                if worksheet.nrows > 0:
                    sheets_content.extend(sheet_content + [''])
                    total_rows += worksheet.nrows
            
            return {
                'content': '\n'.join(sheets_content),
                'metadata': {
                    'sheets': workbook.nsheets,
                    'total_rows': total_rows
                }
            }
            
        except Exception as e:
            return {'content': '', 'error': f'XLS parsing error: {e}'}
    
    def _parse_csv(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse CSV files with automatic delimiter detection."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                # Read sample to detect dialect
                sample = f.read(8192)
                f.seek(0)
                
                # Detect CSV dialect
                try:
                    dialect = csv.Sniffer().sniff(sample)
                    reader = csv.reader(f, dialect)
                except csv.Error:
                    # Fallback to standard CSV
                    f.seek(0)
                    reader = csv.reader(f)
                
                rows = []
                for row_num, row in enumerate(reader):
                    if row_num == 0:
                        # Header row
                        rows.append(' | '.join(f"**{cell}**" for cell in row))
                    else:
                        rows.append(' | '.join(row))
                    
                    # Limit for very large files
                    if row_num > 10000:
                        rows.append(f"... (truncated after {row_num} rows)")
                        break
                
                return {
                    'content': '\n'.join(rows),
                    'metadata': {'rows': len(rows) - 1}  # Exclude header
                }
                
        except Exception as e:
            return {'content': '', 'error': f'CSV parsing error: {e}'}
    
    def _parse_tsv(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse TSV files."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f, delimiter='\t')
                
                rows = []
                for row_num, row in enumerate(reader):
                    if row_num == 0:
                        # Header row
                        rows.append(' | '.join(f"**{cell}**" for cell in row))
                    else:
                        rows.append(' | '.join(row))
                    
                    # Limit for very large files
                    if row_num > 10000:
                        rows.append(f"... (truncated after {row_num} rows)")
                        break
                
                return {
                    'content': '\n'.join(rows),
                    'metadata': {'rows': len(rows) - 1}
                }
                
        except Exception as e:
            return {'content': '', 'error': f'TSV parsing error: {e}'}
    
    def _parse_json(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse JSON and JSONL files."""
        try:
            ext = file_path.suffix.lower()
            
            if ext in {'.jsonl', '.ndjson'}:
                # Line-delimited JSON
                content_lines = []
                with open(file_path, 'r', encoding='utf-8') as f:
                    for line_num, line in enumerate(f):
                        line = line.strip()
                        if line:
                            try:
                                obj = json.loads(line)
                                content_lines.append(json.dumps(obj, indent=2))
                            except json.JSONDecodeError:
                                content_lines.append(f"Invalid JSON on line {line_num + 1}: {line}")
                        
                        # Limit for very large files
                        if line_num > 1000:
                            content_lines.append(f"... (truncated after {line_num} lines)")
                            break
                
                return {
                    'content': '\n---\n'.join(content_lines),
                    'metadata': {'lines': len(content_lines)}
                }
            else:
                # Regular JSON
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Pretty print JSON
                content = json.dumps(data, indent=2, ensure_ascii=False)
                
                return {
                    'content': content,
                    'metadata': {'json_type': type(data).__name__}
                }
                
        except Exception as e:
            return {'content': '', 'error': f'JSON parsing error: {e}'}
    
    def _parse_notebook(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse Jupyter notebook files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                notebook = json.load(f)
            
            content_parts = []
            cell_count = 0
            code_cells = 0
            markdown_cells = 0
            
            for cell in notebook.get('cells', []):
                cell_type = cell.get('cell_type', 'unknown')
                source = cell.get('source', [])
                
                if isinstance(source, list):
                    cell_content = ''.join(source)
                else:
                    cell_content = str(source)
                
                if cell_content.strip():
                    content_parts.append(f"=== {cell_type.title()} Cell ===")
                    content_parts.append(cell_content)
                    content_parts.append('')
                    
                    cell_count += 1
                    if cell_type == 'code':
                        code_cells += 1
                    elif cell_type == 'markdown':
                        markdown_cells += 1
            
            return {
                'content': '\n'.join(content_parts),
                'metadata': {
                    'total_cells': cell_count,
                    'code_cells': code_cells,
                    'markdown_cells': markdown_cells,
                    'kernel': notebook.get('metadata', {}).get('kernelspec', {}).get('name', 'unknown')
                }
            }
            
        except Exception as e:
            return {'content': '', 'error': f'Notebook parsing error: {e}'}
    
    def _parse_html(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse HTML files using BeautifulSoup for clean text extraction."""
        if not HTML_AVAILABLE:
            # Fallback to regex-based parsing
            return self._parse_html_fallback(file_path)
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract title
            title = soup.find('title')
            title_text = title.get_text() if title else ''
            
            # Extract meta description
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            meta_desc_text = meta_desc.get('content', '') if meta_desc else ''
            
            # Extract main content
            # Look for main content areas first
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
            if main_content:
                text = main_content.get_text()
            else:
                text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            # Combine with title and meta
            content_parts = []
            if title_text:
                content_parts.append(f"Title: {title_text}")
            if meta_desc_text:
                content_parts.append(f"Description: {meta_desc_text}")
            if text:
                content_parts.append(text)
            
            final_content = '\n\n'.join(content_parts)
            
            return {
                'content': final_content,
                'metadata': {
                    'title': title_text,
                    'description': meta_desc_text,
                    'links': len(soup.find_all('a')),
                    'images': len(soup.find_all('img')),
                    'forms': len(soup.find_all('form'))
                }
            }
            
        except Exception as e:
            return {'content': '', 'error': f'HTML parsing error: {e}'}
    
    def _parse_html_fallback(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Fallback HTML parsing using regex when BeautifulSoup is not available."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Remove script and style content
            content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
            content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
            
            # Extract title
            title_match = re.search(r'<title[^>]*>(.*?)</title>', content, re.IGNORECASE | re.DOTALL)
            title = title_match.group(1) if title_match else ''
            
            # Remove all HTML tags
            text = re.sub(r'<[^>]+>', ' ', content)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            # Combine with title
            final_content = f"Title: {title}\n\n{text}" if title else text
            
            return {
                'content': final_content,
                'metadata': {'title': title, 'fallback_parser': True}
            }
            
        except Exception as e:
            return {'content': '', 'error': f'HTML fallback parsing error: {e}'}
    
    def _parse_xml(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse XML files including RSS, Atom, and SVG."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            if HTML_AVAILABLE:
                soup = BeautifulSoup(content, 'xml')
                text = soup.get_text()
                
                # Special handling for RSS/Atom feeds
                if soup.find('rss') or soup.find('feed'):
                    items = soup.find_all('item') or soup.find_all('entry')
                    feed_content = []
                    for item in items[:20]:  # Limit to first 20 items
                        title = item.find('title')
                        description = item.find('description') or item.find('summary')
                        if title:
                            feed_content.append(f"Title: {title.get_text()}")
                        if description:
                            feed_content.append(f"Content: {description.get_text()}")
                        feed_content.append("---")
                    
                    text = '\n'.join(feed_content) if feed_content else text
            else:
                # Fallback: remove XML tags
                text = re.sub(r'<[^>]+>', ' ', content)
                text = re.sub(r'\s+', ' ', text).strip()
            
            return {
                'content': text,
                'metadata': {'xml_type': 'rss/atom' if 'rss' in content.lower() or 'feed' in content.lower() else 'xml'}
            }
            
        except Exception as e:
            return {'content': '', 'error': f'XML parsing error: {e}'}
    
    def _parse_email(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse email files (.eml, .msg)."""
        if not EMAIL_AVAILABLE:
            return {
                'content': '',
                'error': 'Email parsing not available. Install with: pip install email-parser'
            }
        
        try:
            with open(file_path, 'rb') as f:
                parser = BytesParser(policy=default)
                msg = parser.parse(f)
            
            # Extract email components
            subject = msg.get('Subject', '')
            sender = msg.get('From', '')
            recipients = msg.get('To', '')
            date = msg.get('Date', '')
            
            # Extract body
            body = ''
            if msg.is_multipart():
                for part in msg.iter_parts():
                    if part.get_content_type() == 'text/plain':
                        body += part.get_content()
                    elif part.get_content_type() == 'text/html':
                        # Extract text from HTML if no plain text found
                        if not body:
                            html_content = part.get_content()
                            if HTML_AVAILABLE:
                                soup = BeautifulSoup(html_content, 'html.parser')
                                body = soup.get_text()
                            else:
                                body = re.sub(r'<[^>]+>', ' ', html_content)
            else:
                body = msg.get_content()
            
            # Combine email parts
            email_parts = []
            if subject:
                email_parts.append(f"Subject: {subject}")
            if sender:
                email_parts.append(f"From: {sender}")
            if recipients:
                email_parts.append(f"To: {recipients}")
            if date:
                email_parts.append(f"Date: {date}")
            if body:
                email_parts.append(f"\nContent:\n{body}")
            
            content = '\n'.join(email_parts)
            
            return {
                'content': content,
                'metadata': {
                    'subject': subject,
                    'sender': sender,
                    'recipients': recipients,
                    'date': date,
                    'attachments': len([part for part in msg.iter_parts() if part.get_content_disposition() == 'attachment']) if msg.is_multipart() else 0
                }
            }
            
        except Exception as e:
            return {'content': '', 'error': f'Email parsing error: {e}'}
    
    def _parse_rtf(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse RTF files using striprtf or pandoc."""
        if PANDOC_AVAILABLE:
            try:
                content = pypandoc.convert_file(str(file_path), 'plain', format='rtf')
                return {
                    'content': content,
                    'metadata': {'parser': 'pandoc'}
                }
            except Exception as e:
                logger.warning(f"Pandoc RTF parsing failed: {e}")
        
        # Fallback: basic RTF parsing
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Basic RTF tag removal
            content = re.sub(r'\\[a-z]+\d*\s*', '', content)  # Remove RTF commands
            content = re.sub(r'[{}]', '', content)  # Remove braces
            content = re.sub(r'\s+', ' ', content).strip()  # Clean whitespace
            
            return {
                'content': content,
                'metadata': {'parser': 'basic_rtf'}
            }
            
        except Exception as e:
            return {'content': '', 'error': f'RTF parsing error: {e}'}
    
    def _parse_archive(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse archive files and extract text from contained files."""
        if not ARCHIVE_AVAILABLE:
            return {
                'content': '',
                'error': 'Archive parsing not available. Install with: pip install zipfile tarfile'
            }
        
        try:
            ext = file_path.suffix.lower()
            extracted_content = []
            file_list = []
            
            if ext == '.zip':
                with zipfile.ZipFile(file_path, 'r') as archive:
                    for file_info in archive.filelist:
                        if not file_info.is_dir():
                            file_list.append(file_info.filename)
                            # Extract text files only (limit size)
                            if file_info.file_size < 1024 * 1024:  # 1MB limit
                                try:
                                    with archive.open(file_info.filename) as f:
                                        content = f.read(8192).decode('utf-8', errors='ignore')
                                        if content.strip():
                                            extracted_content.append(f"=== {file_info.filename} ===")
                                            extracted_content.append(content)
                                except:
                                    continue
                                    
            elif ext in {'.tar', '.tar.gz', '.tgz', '.tar.bz2', '.tbz2', '.tar.xz', '.txz'}:
                with tarfile.open(file_path, 'r:*') as archive:
                    for member in archive.getmembers():
                        if member.isfile():
                            file_list.append(member.name)
                            # Extract text files only (limit size)
                            if member.size < 1024 * 1024:  # 1MB limit
                                try:
                                    f = archive.extractfile(member)
                                    if f:
                                        content = f.read(8192).decode('utf-8', errors='ignore')
                                        if content.strip():
                                            extracted_content.append(f"=== {member.name} ===")
                                            extracted_content.append(content)
                                except:
                                    continue
            
            # Create summary
            summary_parts = [f"Archive contains {len(file_list)} files:"]
            summary_parts.extend([f"  - {name}" for name in file_list[:20]])
            if len(file_list) > 20:
                summary_parts.append(f"  ... and {len(file_list) - 20} more files")
            
            if extracted_content:
                summary_parts.extend(["", "=== Extracted Text Content ==="] + extracted_content)
            
            content = '\n'.join(summary_parts)
            
            return {
                'content': content,
                'metadata': {
                    'total_files': len(file_list),
                    'extracted_files': len(extracted_content) // 2,  # Each file has header + content
                    'archive_type': ext
                }
            }
            
        except Exception as e:
            return {'content': '', 'error': f'Archive parsing error: {e}'}
    
    def _parse_text(self, file_path: Path) -> Dict[str, Union[str, Dict]]:
        """Parse text and code files with encoding detection."""
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()
                
                # Clean and normalize content
                content = self._clean_text(content)
                
                return {
                    'content': content,
                    'metadata': {
                        'encoding': encoding,
                        'lines': content.count('\n') + 1 if content else 0,
                        'characters': len(content)
                    }
                }
                
            except UnicodeDecodeError:
                continue
            except Exception as e:
                return {'content': '', 'error': f'Text parsing error: {e}'}
        
        return {'content': '', 'error': 'Could not decode file with any supported encoding'}
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ''
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive blank lines
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        return text.strip()

# Global parser instance
parser = FileParser()

def parse_file(file_path: Union[str, Path]) -> Dict[str, Union[str, Dict]]:
    """
    Convenience function to parse a file.
    
    Args:
        file_path: Path to file to parse
        
    Returns:
        Dictionary with parsed content and metadata
    """
    return parser.parse_file(file_path)

def get_supported_extensions() -> set:
    """Get all supported file extensions."""
    return FileParser.get_supported_extensions()

def is_supported_file(file_path: Union[str, Path]) -> bool:
    """Check if file type is supported."""
    return FileParser.is_supported(file_path) 